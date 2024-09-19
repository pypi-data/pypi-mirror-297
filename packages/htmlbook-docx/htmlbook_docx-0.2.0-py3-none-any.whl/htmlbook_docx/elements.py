import logging
from typing import Union
from bs4 import Tag, NavigableString
from docx import Document
from docx.shared import Inches
from docx.styles.style import CharacterStyle, ParagraphStyle
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK_TYPE


def process_tag(doc: Document, element: Tag):
    """
    Case-processes tags
    """

    match element.name:
        case "h1":
            doc.add_heading(element.string)
        case "p":
            add_paragraph_from_tag(doc, element)
        case "div":
            for ele in element.contents:
                if isinstance(ele, Tag):
                    match ele.name:
                        case "p":
                            # pass parent style through
                            add_paragraph_from_tag(
                                doc, ele, get_style(doc, element)
                            )
                        case "dl":
                            add_definiton_list(doc, ele)
        case "hr":
            break_p = doc.add_paragraph("#")
            break_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        case "blockquote":  # songs
            if element.pre:
                for piece in element.pre:
                    if isinstance(piece, NavigableString):
                        paras = piece.split("\n")
                        for para in paras:
                            doc.add_paragraph(para).style.name = "verse"
                    else:
                        print(f"confused about {piece}")
            else:
                process_tag(doc, element)

        case "dl":
            add_definiton_list(doc, element)

        case _:
            logging.warning(f"Unexpected or unhandled tag: {element}")


def add_paragraph_from_tag(doc: Document, p_tag: Tag, style=None):
    """
    Handles adding a paragraph, including inlines and runs, to the document
    """
    p = doc.add_paragraph()

    # overrides and passes
    if p_tag.attrs.get("class", ""):
        p.style = get_style(doc, p_tag)
        style_pass = p.style.name
    elif style:
        p.style = style
        style_pass = style.name
    else:
        style_pass = ""
        p.style.name = "Normal"

    # loop through paragraph components
    for ele in p_tag.contents:
        match ele:
            case NavigableString():
                p.add_run(sanitize(ele))
            case Tag():
                add_runs(doc, p, ele, style_pass)  # type: ignore
            case _:
                logging.warning(f"Missed content: {ele}")

    return p


def add_runs(
    doc: Document,
    p: Paragraph,
    inline: Tag,
    parent_style="",
    combine_newlines=True,
):
    """
    Adds a run to a given paragraph from an inline tag; if an inline tag itself
    contains further inline tags, those will be added as separate runs
    """

    for ele in inline.contents:
        match ele:
            case NavigableString():
                # clears away the unwanted line breaks
                if combine_newlines:
                    ele = sanitize(ele)
                match inline.name:
                    case "em" | "i":
                        if parent_style in ["instruction", "playintro"]:
                            p.add_run(f"{ele}").italic = False
                        else:
                            p.add_run(f"{ele}").italic = True
                    case "strong" | "b":
                        p.add_run(f"{ele}").bold = True
                    case "code":
                        # style as small caps for now
                        run = p.add_run(f"{ele}")
                        run.font.small_caps = True
                        if parent_style in ["instruction", "playintro"]:
                            run.italic = False
                    case "span":
                        # TODO: handle classes better
                        run = p.add_run(f"{ele}", style=get_style(doc, inline))
                    case _:  # just add the text
                        p.add_run(f"{ele}")
            case Tag():
                if ele.name == "br":
                    p.add_run().add_break(WD_BREAK_TYPE.LINE)
                else:
                    add_runs(doc, p, ele)
            case _:
                logging.warning(f"Missed content: {ele}")



def add_definiton_list(doc: Document, dl: Tag):
    """
    Adds a definition list... sort of
    """
    dialogue_p = None
    for tag in dl.contents:
        if not isinstance(tag, Tag):
            logging.info("Non-tag in dl: {tag}")
            pass
        elif tag.name == "dt":
            dialogue_p = doc.add_paragraph()
            dialogue_p.style = doc.styles["playspace"]
            dialogue_p.paragraph_format.left_indent = Inches(0.5)
            dialogue_p.paragraph_format.first_line_indent = Inches(-0.5)
            if tag.string:
                tagstring = tag.string.replace("\n", "")
                dialogue_p.add_run(f"{tagstring}:\t").font.small_caps = True
            else:
                raise Exception
        elif tag.name == "dd":
            if tag.string:  # unlikely
                tagstring = tag.string.replace("\n", "")
                dialogue_p.add_run(tagstring)  # type: ignore
            else:
                for ele in tag.contents:
                    match ele.name:  # type:ignore
                        case "p":
                            add_runs(doc, dialogue_p, ele)  # type: ignore
                        case "blockquote":
                            if ele.pre:  # type: ignore
                                add_runs(
                                    doc,
                                    dialogue_p,  # type: ignore
                                    ele.pre,  # type: ignore
                                    combine_newlines=False,
                                )
                                dialogue_p.paragraph_format.line_spacing = 1.0  # type: ignore
                        case "div":  # open block!
                            handle_div_in_dl(doc, dialogue_p, ele)

                        case _:
                            try:
                                if not ele.strip() == "":  # type: ignore
                                    logging.warning(f"Missed content: {ele}")
                            except Exception:
                                logging.warning(f"Missed content: {ele}")


def handle_div_in_dl(doc, dialogue_p, div):
    added = False
    for p in div.contents:  # type: ignore
        if str(p).strip() != "":
            if not added:
                add_runs(doc, dialogue_p, p)  # type: ignore
                added = True
            else:
                if isinstance(p, NavigableString):
                    add_runs(doc, dialogue_p, p)  # type: ignore
                else:  # make a new p
                    dialogue_p = add_paragraph_from_tag(
                        doc, p, doc.styles["playspace"]
                    )
                    dialogue_p.paragraph_format.left_indent = Inches(
                        0.5
                    )
                    dialogue_p.paragraph_format.first_line_indent = Inches(
                        0
                    )

def get_style(doc: Document, tag: Tag) -> Union[ParagraphStyle, CharacterStyle]:
    style_name = tag.attrs.get("class", [])[-1]  # "openblock aside"
    style_name = style_name.replace("-", "")
    if style_name == "ii":
        style_name = "inlineinstruction"
    return doc.styles[style_name]


def sanitize(s: NavigableString) -> str:
    return str(s).replace("\n", " ")
