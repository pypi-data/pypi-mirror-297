#import libs
from __future__ import print_function
import os
import re
import PyPDF2
import fitz
import os.path
import io
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.http import MediaFileUpload
from docx import Document
import configparser
import nltk
#from nltk.corpus import stopwords
#from nltk.tokenize import sent_tokenize, word_tokenize
#from nltk.stem import WordNetLemmatizer
import string
import openai
from langchain_openai import ChatOpenAI
import openai
from groq import Groq
from langchain.chains import LLMChain, RetrievalQA
#import time
#import re
import warnings
from langchain.memory import ConversationBufferMemory
from langchain.schema import HumanMessage
from langchain.prompts import ChatPromptTemplate
from langchain.chains import ConversationChain
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.runnables.base import Runnable
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
import uuid
from datetime import datetime, timedelta
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import WordNetLemmatizer
import string
from langchain.embeddings import HuggingFaceInstructEmbeddings
#from InstructorEmbedding import INSTRUCTOR
from sklearn.cluster import KMeans
import numpy as np
import voyageai
from langchain_voyageai import VoyageAIEmbeddings
from sklearn.metrics.pairwise import cosine_similarity
from rouge_score import rouge_scorer
from ipywidgets import widgets
from IPython.display import display


nltk.download('punkt')
nltk.download('wordnet')

class GoogleDriveHandler:
    def __init__(self, credentials_path='credentials.json', token_path='token.json'):
        self.SCOPES = ['https://www.googleapis.com/auth/drive']
        self.credentials_path = credentials_path
        self.token_path = token_path
        self.service = self.create_service()

    def create_service(self):
        creds = None

        if not os.path.exists(self.credentials_path):
            raise FileNotFoundError(f"{self.credentials_path} not found. Please ensure it is in your current working directory.")

        if os.path.exists(self.token_path):
            try:
                creds = Credentials.from_authorized_user_file(self.token_path, self.SCOPES)
            except Exception as e:
                print(f"Error loading {self.token_path}: {e}")
                creds = None

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    creds.refresh(Request())
                except Exception as e:
                    print(f"Failed to refresh token: {e}. Re-authenticating...")
                    creds = None

            if not creds:
                try:
                    flow = InstalledAppFlow.from_client_secrets_file(self.credentials_path, self.SCOPES)
                    creds = flow.run_local_server(port=0)
                except Exception as e:
                    raise RuntimeError(f"Failed to obtain credentials: {e}")

            with open(self.token_path, 'w') as token:
                token.write(creds.to_json())

        try:
            service = build('drive', 'v3', credentials=creds)
        except Exception as e:
            raise RuntimeError(f"Failed to create the Google Drive service: {e}")

        return service

    def ensure_directory(self, path):
        if not os.path.exists(path):
            os.makedirs(path)

    def get_files_in_folder(self, folder_id, mimeType, page_size=10):
        query = f"'{folder_id}' in parents and mimeType='{mimeType}' and trashed=false"
        files = []
        page_token = None
        while True:
            results = self.service.files().list(
                q=query,
                spaces='drive',
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=page_size
            ).execute()
            files.extend(results.get('files', []))
            page_token = results.get('nextPageToken', None)
            if page_token is None:
                break
        return files

    def get_files_in_folder_with_query(self, query, page_size=10):
        query += " and trashed=false"
        files = []
        page_token = None
        while True:
            results = self.service.files().list(
                q=query,
                spaces='drive',
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=page_size
            ).execute()
            files.extend(results.get('files', []))
            page_token = results.get('nextPageToken', None)
            if page_token is None:
                break
        return files

    def download_file(self, item, save_dir):
        file_name = item['name']
        if not os.path.exists(os.path.join(save_dir, file_name)):
            print(f"Downloading {file_name}...")
            request = self.service.files().get_media(fileId=item['id'])
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
                print(f"Download {int(status.progress() * 100)}%.")
            with open(os.path.join(save_dir, file_name), 'wb') as f:
                fh.seek(0)
                f.write(fh.read())
        else:
            print(f"{file_name} already exists. Skipping download.")

    def get_existing_files(self, folder_id):
        existing_files = self.service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            spaces='drive',
            fields='nextPageToken, files(id, name)'
        ).execute()
        return [file['name'] for file in existing_files.get('files', [])]

    def upload_file(self, file_name, folder_id, directory_path):
        file_path = os.path.join(directory_path, file_name)
        file_metadata = {'name': file_name, 'parents': [folder_id]}
        media = MediaFileUpload(file_path, mimetype='text/plain')
        file = self.service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        print(f"{file_name} uploaded successfully with File ID: {file.get('id')}")

    def download_pdfs(self, folder_id, save_dir='PDF_docs', limit = None):
        self.ensure_directory(save_dir)
        page_token = None
        downloaded_files = 0
        while True:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and mimeType='application/pdf' and trashed=false",
                spaces='drive',
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=10
            ).execute()
            items = results.get('files', [])
            if not items:
                print('No more files found.')
                break
            for item in items:
                if limit is not None and downloaded_files >= limit:
                    print(f"Download limit of {limit} files reached.")
                    return
                self.download_file(item, save_dir)
                downloaded_files += 1
            page_token = results.get('nextPageToken', None)
            if page_token is None:
                break

    def upload_txt(self, folder_id, directory_path='.'):
        self.ensure_directory(directory_path)
        files = [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f)) and f.endswith('.txt')]
        existing_files = self.get_existing_files(folder_id)

        for file_name in files:
            if file_name not in existing_files:
                self.upload_file(file_name, folder_id, directory_path)

    def convert_pdf_to_text(self, pdf_path):
        text = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        return text

    def process_pdfs_in_dir(self, directory_path):
        self.ensure_directory(directory_path)
        for filename in os.listdir(directory_path):
            if filename.lower().endswith('.pdf'):
                full_path = os.path.join(directory_path, filename)
                pdf_text = self.convert_pdf_to_text(full_path)
                output_filename = filename.rsplit('.', 1)[0] + '.txt'
                output_path = os.path.join(directory_path, output_filename)
                with open(output_path, 'w', encoding='utf-8') as text_file:
                    text_file.write(pdf_text)
                print(f"Processed and saved: {filename} as {output_filename}")

    def docx_to_text(self, docx_file_path):
        doc = Document(docx_file_path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(text)

    def convert_docx_to_txt(self, folder_path):
        self.ensure_directory(folder_path)
        for file_name in os.listdir(folder_path):
            if file_name.endswith('.docx'):
                docx_file_path = os.path.join(folder_path, file_name)
                text_file_name = file_name.replace('.docx', '.txt')
                text_file_path = os.path.join(folder_path, text_file_name)
                text_content = self.docx_to_text(docx_file_path)
                with open(text_file_path, 'w', encoding='utf-8') as text_file:
                    text_file.write(text_content)
                print(f"Converted {file_name} to {text_file_name}")

    def download_txt(self, folder_id, save_dir='Text_docs', limit = None):
        self.ensure_directory(save_dir)
        page_token = None
        downloaded_files = 0
        
        while True:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and mimeType='text/plain' and trashed=false",
                spaces='drive',
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=10
            ).execute()
            items = results.get('files', [])
            if not items:
                print('No more files found.')
                break
            for item in items:
                if limit is not None and downloaded_files >= limit:
                    print(f"Download limit of {limit} files reached.")
                    return
                self.download_file(item, save_dir)
                downloaded_files += 1
            page_token = results.get('nextPageToken', None)
            if page_token is None:
                break

    def download_docs(self, folder_id, save_dir='Doc_docs', limit = None):
        self.ensure_directory(save_dir)
        query = f"'{folder_id}' in parents and (mimeType='application/msword' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') and trashed=false"
        page_token = None
        downloaded_files = 0
        
        while True:
            results = self.service.files().list(
                q=query,
                spaces='drive',
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=10
            ).execute()
            items = results.get('files', [])
            if not items:
                print('No more files found.')
                break
            for item in items:
                if limit is not None and downloaded_files >= limit:
                    print(f"Download limit of {limit} files reached.")
                    return
                self.download_file(item, save_dir)
                downloaded_files += 1
            page_token = results.get('nextPageToken', None)
            if page_token is None:
                break
    
    def upload_docs(self, folder_id, directory_path='.'):
        self.ensure_directory(directory_path)
        files = [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f)) and (f.endswith('.docx') or f.endswith('.doc'))]
        existing_files = self.get_existing_files(folder_id)

        for file_name in files:
            if file_name not in existing_files:
                self.upload_file(file_name, folder_id, directory_path)    
      
      
      #This part add LLM to the package allowing users to summarize PDFs easily
      #Begin by pre-processing the data
      
    def preprocess_text(self, text):
        lemmatizer = WordNetLemmatizer()
        sentences = nltk.sent_tokenize(text)
        punctuation = set(string.punctuation)

        processed_sentences = []
        for sent in sentences:
            words = nltk.word_tokenize(sent)
            filtered_words = [
                lemmatizer.lemmatize(word.lower()) 
                for word in words 
                if word.lower() not in punctuation and word.isalpha()
            ]
            processed_sentences.append(' '.join(filtered_words))

        processed_text = ' '.join(processed_sentences)
        processed_text = re.sub(r'\d+', '', processed_text)

        return processed_text

    def extract_text_from_pdf(self, pdf_path):
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def extract_sections(self, text):
        sections = {
            "methodology": "",
            "methods": "",
            "results": "",
            "discussion": "",
            "conclusion": ""
        }
    
        current_section = None
        start_extracting = False
        is_discussion = False
    
        for line in text.split('\n'):
            line_lower = line.strip().lower()

            # Stop extracting when "references" section is encountered
            if line_lower.startswith("references"):
                start_extracting = False
        
            # Start extracting from "methodology" sections
            elif ("methodology" in line_lower or 
                "methods" in line_lower or 
                "materials and methods" in line_lower or 
                "materials & methods" in line_lower) and not is_discussion:
                
                current_section = "methodology"
                start_extracting = True
        
            # Start extracting from the "results" section
            elif "results" in line_lower and not is_discussion:
                current_section = "results"
                start_extracting = True
        
            # Start extracting from the "discussion" section
            elif "discussion" in line_lower:
                current_section = "discussion"
                is_discussion = True
                start_extracting = True
        
            # Start extracting from the "conclusion" section
            elif "conclusion" in line_lower:
                current_section = "conclusion"
                start_extracting = True
        
            # Stop extracting when "acknowledgements" section is encountered
            elif "acknowledgements" in line_lower:
                start_extracting = False
        
            # Add lines to the current section if extracting is active
            if start_extracting and current_section:
                sections[current_section] += line + "\n"

        # Combine the extracted sections
        combined_text = (sections["methodology"] + sections["results"] + 
                         sections["discussion"] + sections["conclusion"])
    
        return combined_text, sections


    def get_model(self, selected_model, OPENAI_API_KEY, GROQ_API_KEY):
        if selected_model == "llama3-8b-8192":
            return ChatGroq(groq_api_key=GROQ_API_KEY, model="llama3-8b-8192", temperature=0.02, max_tokens=None, timeout=None, max_retries=2)
        elif selected_model == "llama3-70b-8192":
            return ChatGroq(groq_api_key=GROQ_API_KEY, model="llama3-70b-8192", temperature=0.02, max_tokens=None, timeout=None, max_retries=2) 
        elif selected_model == "gpt-4o-mini":
            return ChatOpenAI(model="gpt-4o-mini", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        elif selected_model == "gpt-4o":
            return ChatOpenAI(model="gpt-4o", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        elif selected_model == "gpt-4":
            return ChatOpenAI(model="gpt-4", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        else:
            raise ValueError("Invalid model selected")

    def chunk_text_with_langchain(self, text, chunk_size=8000, chunk_overlap=500):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_text(text)
        return chunks

    def embed_chunks(self, chunks, VOYAGEAI_API_key):
        vo = voyageai.Client(api_key=VOYAGEAI_API_key)
        result = vo.embed(chunks, model="voyage-large-2-instruct", input_type="document")
        vectors = result.embeddings
        return np.array(vectors)

    def clustering(self, vectors, num_clusters):
        kmeans = KMeans(n_clusters=num_clusters, random_state=42).fit(vectors)
        labels = kmeans.labels_

        closest_indices = []
        for i in range(num_clusters):
            distances = np.linalg.norm(vectors - kmeans.cluster_centers_[i], axis=1)
            closest_index = np.argmin(distances)
            closest_indices.append(closest_index)

        selected_indices = sorted(closest_indices)
        return selected_indices

    def filter_redundant_chunks(self, chunks, vectors, similarity_threshold=0.8):
        unique_chunks = []
        unique_vectors = []

        for i, vector in enumerate(vectors):
            if len(unique_vectors) == 0:
                unique_chunks.append(chunks[i])
                unique_vectors.append(vector)
            else:
                similarities = cosine_similarity([vector], unique_vectors)
                if max(similarities[0]) < similarity_threshold:
                    unique_chunks.append(chunks[i])
                    unique_vectors.append(vector)

        return unique_chunks, unique_vectors
    
    
    def summarize_text(self, text, selected_model, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size=8000, chunk_overlap=500, similarity_threshold=0.8, num_clusters=10):
        llm_mod = self.get_model(selected_model, OPENAI_API_KEY, GROQ_API_KEY)
        system_prompt = prompt
        prompt_template = ChatPromptTemplate.from_messages([
        SystemMessage(content=system_prompt),
        HumanMessagePromptTemplate.from_template("{text}")
        ])
        
        conversation = LLMChain(llm=llm_mod, prompt=prompt_template)
        
        if selected_model in ["llama3-8b-8192", "llama3-70b-8192", "gpt-4"]:
            chunks = self.chunk_text_with_langchain(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            vectors = self.embed_chunks(chunks, VOYAGEAI_API_key)
            
            unique_chunks, unique_vectors = self.filter_redundant_chunks(chunks, vectors, similarity_threshold=similarity_threshold)
            
            num_clusters = min(num_clusters, len(unique_chunks))
            
            selected_indices = self.clustering(unique_vectors, num_clusters)
            selected_chunks = [unique_chunks[i] for i in selected_indices]
            selected_text = ' '.join(selected_chunks)
            
            if len(selected_text) > chunk_size:
                final_summary_chunks = []
                for i in range(0, len(selected_text), chunk_size):
                    final_summary_chunks.append(conversation.run(selected_text[i:i + chunk_size]))
                summary = ' '.join(final_summary_chunks)
            
            else:
                summary = conversation.run(selected_text)
        else:
            summary = conversation.run(text)
        return summary

    
    def save_summary_as_docx(self, summary, output_path, pdf_filename):
        doc = Document()
        title = os.path.splitext(pdf_filename)[0]
        doc.add_heading(f'{title} Summary', 0)
        doc.add_paragraph(summary)
        doc.save(output_path)

    

    def summarize_pdfs(self, pdf_directory, output_directory, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size=8000, chunk_overlap=500, similarity_threshold=0.8, num_clusters=10):
        model_options = ["llama3-8b-8192", "llama3-70b-8192", "gpt-4o-mini", "gpt-4o", "gpt-4"]
    
        model_dropdown = widgets.Dropdown(
            options=model_options,
            value=model_options[0],
            description='Select Model:',
        )
    
        display(model_dropdown)
    
        status_label = widgets.Label(value="Select a model to start processing...")
        display(status_label)
    
        progress_bar = widgets.IntProgress(
            value=0,
            min=0,
            max=100,
            step=1,
            description='Progress:',
            bar_style='info',
            orientation='horizontal'
        )
        display(progress_bar)
    
        process_button = widgets.Button(description="Start Processing")
        display(process_button)
    
        def on_button_click(b):
            selected_model = model_dropdown.value
            status_label.value = "Processing... Please wait."
        
            pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('.pdf')]
            total_files = len(pdf_files)
        
            for idx, pdf_filename in enumerate(pdf_files):
                pdf_path = os.path.join(pdf_directory, pdf_filename)
                output_path = os.path.join(output_directory, f"Summary-{os.path.splitext(pdf_filename)[0]}.docx")
            
                # Check if the summary already exists
                if os.path.exists(output_path):
                    print(f"Summary already exists for {pdf_filename}. Skipping...")
                    continue
            
                text = self.extract_text_from_pdf(pdf_path)
            
                # Skip processing if extracted text is empty
                if not text.strip():
                    print(f"No text found in {pdf_filename}. Skipping...")
                    continue
            
                combined_text, _ = self.extract_sections(text)
                preprocessed_text = self.preprocess_text(combined_text)
            
                # Skip processing if preprocessed text is empty
                if not preprocessed_text.strip():
                    print(f"No meaningful text after preprocessing for {pdf_filename}. Skipping...")
                    continue
            
                # Chunking the preprocessed text
                chunks = self.chunk_text_with_langchain(preprocessed_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            
                # Skip embedding if there are no chunks
                if not chunks:
                    print(f"No chunks generated for {pdf_filename}. Skipping...")
                    continue
            
                vectors = self.embed_chunks(chunks, VOYAGEAI_API_key)
            
                # Continue with the summarization if chunks and vectors are valid
                summary = self.summarize_text(preprocessed_text, selected_model, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size, chunk_overlap, similarity_threshold, num_clusters)
            
                self.save_summary_as_docx(summary, output_path, pdf_filename)
            
                # Update progress bar
                progress = int((idx + 1) / total_files * 100)
                progress_bar.value = progress
                progress_bar.description = f'Progress: {progress}%'
        
        status_label.value = "Processing complete. Summaries saved."
        progress_bar.bar_style = 'success'
        progress_bar.description = 'Complete'
    
        process_button.on_click(on_button_click)


# Entry point for command line usage
def main():
    import argparse

    parser = argparse.ArgumentParser(description='Google Drive Handler')
    parser.add_argument('action', choices=['download_pdfs', 'upload_txt', 'convert_pdfs', 'convert_docx', 'download_txts', 'download_docs', 'run_app'], help='Action to perform')
    parser.add_argument('folder_id', help='Google Drive folder ID')
    parser.add_argument('--credentials', default='credentials.json', help='Path to credentials.json')
    parser.add_argument('--directory', default='.', help='Directory to process files in')
    parser.add_argument('--model', default='gpt-4', help='Model to use for summarization')
    parser.add_argument('--output', default='summary_folder', help='Output directory for summaries')

    args = parser.parse_args()

    handler = GoogleDriveHandler(credentials_path=args.credentials)
    
    
    if args.action == 'download_pdfs':
        handler.download_pdfs(args.folder_id)
    elif args.action == 'upload_txt':
        handler.upload_txt(args.folder_id, directory_path=args.directory)
    elif args.action == 'convert_pdfs':
        handler.process_pdfs_in_dir(args.directory)
    elif args.action == 'convert_docx':
        handler.convert_docx_to_txt(args.directory)
    elif args.action == 'download_txts':
        handler.download_txt(args.folder_id, save_dir=args.directory)
    elif args.action == 'download_docs':
        handler.download_docs(args.folder_id, save_dir=args.directory)
    elif args.action == 'upload_docs':
        handler.upload_docs(args.folder_id, directory_path=args.directory)
    elif args.action == 'summaerize_pdfs':
        handler.summarize_pdfs(args.directory, args.output, args.model, os.getenv("My_OpenAI_API_key"), os.getenv("My_Groq_API_key"), os.getenv("My_voyageai_API_key"))
        

if __name__ == '__main__':
    main()
