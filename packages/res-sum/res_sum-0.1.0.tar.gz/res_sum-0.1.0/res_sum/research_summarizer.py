from __future__ import print_function
import os
import re
import PyPDF2
import fitz
import os.path
import io
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.http import MediaFileUpload
from docx import Document
import configparser
from GDriveOps.GDhandler import GoogleDriveHandler
import nltk
import string
import openai
from langchain_openai import ChatOpenAI
import openai
from groq import Groq
from langchain.chains import LLMChain, RetrievalQA
import warnings
from langchain.memory import ConversationBufferMemory
from langchain.schema import HumanMessage
from langchain.prompts import ChatPromptTemplate
from langchain.chains import ConversationChain
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.runnables.base import Runnable
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
import uuid
from datetime import datetime, timedelta
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import WordNetLemmatizer
import string
from langchain.embeddings import HuggingFaceInstructEmbeddings
from sklearn.cluster import KMeans
import numpy as np
import voyageai
from langchain_voyageai import VoyageAIEmbeddings
from sklearn.metrics.pairwise import cosine_similarity
from rouge_score import rouge_scorer
from ipywidgets import widgets
from IPython.display import display


nltk.download('punkt_tab')
#nltk.download('punkt')
#nltk.download('stopwords')
#nltk.download('wordnet')


class ResSum:
    def preprocess_text(self, text):
        lemmatizer = WordNetLemmatizer()
        sentences = nltk.sent_tokenize(text)
        punctuation = set(string.punctuation)

        processed_sentences = []
        for sent in sentences:
            words = nltk.word_tokenize(sent)
            filtered_words = [
                lemmatizer.lemmatize(word.lower()) 
                for word in words 
                if word.lower() not in punctuation and word.isalpha()
            ]
            processed_sentences.append(' '.join(filtered_words))

        processed_text = ' '.join(processed_sentences)
        processed_text = re.sub(r'\d+', '', processed_text)

        return processed_text

    def extract_text_from_pdf(self, pdf_path):
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def extract_sections(self, text):
        sections = {
            "methodology": "",
            "methods": "",
            "results": "",
            "discussion": "",
            "conclusion": ""
        }
    
        current_section = None
        start_extracting = False
        is_discussion = False
    
        for line in text.split('\n'):
            line_lower = line.strip().lower()

            # Stop extracting when "references" section is encountered
            if line_lower.startswith("references"):
                start_extracting = False
        
            # Start extracting from "methodology" sections
            elif ("methodology" in line_lower or 
                "methods" in line_lower or 
                "materials and methods" in line_lower or 
                "materials & methods" in line_lower) and not is_discussion:
                
                current_section = "methodology"
                start_extracting = True
        
            # Start extracting from the "results" section
            elif "results" in line_lower and not is_discussion:
                current_section = "results"
                start_extracting = True
        
            # Start extracting from the "discussion" section
            elif "discussion" in line_lower:
                current_section = "discussion"
                is_discussion = True
                start_extracting = True
        
            # Start extracting from the "conclusion" section
            elif "conclusion" in line_lower:
                current_section = "conclusion"
                start_extracting = True
        
            # Stop extracting when "acknowledgements" section is encountered
            elif "acknowledgements" in line_lower:
                start_extracting = False
        
            # Add lines to the current section if extracting is active
            if start_extracting and current_section:
                sections[current_section] += line + "\n"

        # Combine the extracted sections
        combined_text = (sections["methodology"] + sections["results"] + 
                         sections["discussion"] + sections["conclusion"])
    
        return combined_text, sections


    def get_model(self, selected_model, OPENAI_API_KEY, GROQ_API_KEY):
        if selected_model == "llama3-8b-8192":
            return ChatGroq(groq_api_key=GROQ_API_KEY, model="llama3-8b-8192", temperature=0.02, max_tokens=None, timeout=None, max_retries=2)
        elif selected_model == "llama3-70b-8192":
            return ChatGroq(groq_api_key=GROQ_API_KEY, model="llama3-70b-8192", temperature=0.02, max_tokens=None, timeout=None, max_retries=2) 
        elif selected_model == "gpt-4o-mini":
            return ChatOpenAI(model="gpt-4o-mini", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        elif selected_model == "gpt-4o":
            return ChatOpenAI(model="gpt-4o", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        elif selected_model == "gpt-4":
            return ChatOpenAI(model="gpt-4", temperature=0, max_tokens=None, timeout=None, max_retries=2, api_key=OPENAI_API_KEY)
        else:
            raise ValueError("Invalid model selected")

    def chunk_text_with_langchain(self, text, chunk_size=8000, chunk_overlap=500):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_text(text)
        return chunks

    def embed_chunks(self, chunks, VOYAGEAI_API_key):
        vo = voyageai.Client(api_key=VOYAGEAI_API_key)
        result = vo.embed(chunks, model="voyage-large-2-instruct", input_type="document")
        vectors = result.embeddings
        return np.array(vectors)

    def clustering(self, vectors, num_clusters):
        kmeans = KMeans(n_clusters=num_clusters, random_state=42).fit(vectors)
        labels = kmeans.labels_

        closest_indices = []
        for i in range(num_clusters):
            distances = np.linalg.norm(vectors - kmeans.cluster_centers_[i], axis=1)
            closest_index = np.argmin(distances)
            closest_indices.append(closest_index)

        selected_indices = sorted(closest_indices)
        return selected_indices

    def filter_redundant_chunks(self, chunks, vectors, similarity_threshold=0.8):
        unique_chunks = []
        unique_vectors = []

        for i, vector in enumerate(vectors):
            if len(unique_vectors) == 0:
                unique_chunks.append(chunks[i])
                unique_vectors.append(vector)
            else:
                similarities = cosine_similarity([vector], unique_vectors)
                if max(similarities[0]) < similarity_threshold:
                    unique_chunks.append(chunks[i])
                    unique_vectors.append(vector)

        return unique_chunks, unique_vectors
    
    
    def summarize_text(self, text, selected_model, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size=8000, chunk_overlap=500, similarity_threshold=0.8, num_clusters=10):
        llm_mod = self.get_model(selected_model, OPENAI_API_KEY, GROQ_API_KEY)
        system_prompt = prompt
        prompt_template = ChatPromptTemplate.from_messages([
        SystemMessage(content=system_prompt),
        HumanMessagePromptTemplate.from_template("{text}")
        ])
        
        conversation = LLMChain(llm=llm_mod, prompt=prompt_template)
        
        if selected_model in ["llama3-8b-8192", "llama3-70b-8192", "gpt-4"]:
            chunks = self.chunk_text_with_langchain(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            vectors = self.embed_chunks(chunks, VOYAGEAI_API_key)
            
            unique_chunks, unique_vectors = self.filter_redundant_chunks(chunks, vectors, similarity_threshold=similarity_threshold)
            
            num_clusters = min(num_clusters, len(unique_chunks))
            
            selected_indices = self.clustering(unique_vectors, num_clusters)
            selected_chunks = [unique_chunks[i] for i in selected_indices]
            selected_text = ' '.join(selected_chunks)
            
            if len(selected_text) > chunk_size:
                final_summary_chunks = []
                for i in range(0, len(selected_text), chunk_size):
                    final_summary_chunks.append(conversation.run(selected_text[i:i + chunk_size]))
                summary = ' '.join(final_summary_chunks)
            
            else:
                summary = conversation.run(selected_text)
        else:
            summary = conversation.run(text)
        return summary

    #save summary as docs
    
    def save_summary_as_docx(self, summary, output_path, pdf_filename):
        doc = Document()
        title = f'Summary- {os.path.splitext(pdf_filename)[0]}'
        doc.add_heading(title, 0)
        doc.add_paragraph(summary)
        doc.save(output_path)

    

    def summarize_pdfs(self, pdf_directory, output_directory, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size=8000, chunk_overlap=500, similarity_threshold=0.8, num_clusters=10):
        model_options = ["llama3-8b-8192", "llama3-70b-8192", "gpt-4o-mini", "gpt-4o", "gpt-4"]
    
        model_dropdown = widgets.Dropdown(
            options=model_options,
            value=model_options[0],
            description='Select Model:',
        )
    
        display(model_dropdown)
    
        status_label = widgets.Label(value="Select a model to start processing...")
        display(status_label)
    
        progress_bar = widgets.IntProgress(
            value=0,
            min=0,
            max=100,
            step=1,
            description='Progress:',
            bar_style='info',
            orientation='horizontal'
        )
        display(progress_bar)
    
        process_button = widgets.Button(description="Start Processing")
        display(process_button)
    
        def on_button_click(b):
            selected_model = model_dropdown.value
            status_label.value = "Processing... Please wait."
        
            pdf_files = [f for f in os.listdir(pdf_directory) if f.endswith('.pdf')]
            total_files = len(pdf_files)
        
            for idx, pdf_filename in enumerate(pdf_files):
                pdf_path = os.path.join(pdf_directory, pdf_filename)
                output_path = os.path.join(output_directory, f"Summary-{os.path.splitext(pdf_filename)[0]}.docx")
            
                #print(f"Processing PDF: {pdf_path}")  # Debug
                #print(f"Output path: {output_path}")  # Debug 
                
                # Check if the summary already exists
                if os.path.exists(output_path):
                    print(f"Summary already exists for {pdf_filename}. Skipping...")
                    continue
            
                text = self.extract_text_from_pdf(pdf_path)
            
                # Skip processing if extracted text is empty
                if not text.strip():
                    print(f"No text found in {pdf_filename}. Skipping...")
                    continue
            
                combined_text, _ = self.extract_sections(text)
                preprocessed_text = self.preprocess_text(combined_text)
            
                # Skip processing if preprocessed text is empty
                if not preprocessed_text.strip():
                    print(f"No meaningful text after preprocessing for {pdf_filename}. Skipping...")
                    continue
            
                # Chunking the preprocessed text
                chunks = self.chunk_text_with_langchain(preprocessed_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            
                # Skip embedding if there are no chunks
                if not chunks:
                    print(f"No chunks generated for {pdf_filename}. Skipping...")
                    continue
            
                vectors = self.embed_chunks(chunks, VOYAGEAI_API_key)
            
                # Continue with the summarization if chunks and vectors are valid
                summary = self.summarize_text(preprocessed_text, selected_model, prompt, OPENAI_API_KEY, GROQ_API_KEY, VOYAGEAI_API_key, chunk_size, chunk_overlap, similarity_threshold, num_clusters)
            
                self.save_summary_as_docx(summary, output_path, pdf_filename)
            
                # Update progress bar
                progress = int((idx + 1) / total_files * 100)
                progress_bar.value = progress
                progress_bar.description = f'Progress: {progress}%'
        
        status_label.value = "Processing complete. Summaries saved."
        progress_bar.bar_style = 'success'
        progress_bar.description = 'Complete'
    
        process_button.on_click(on_button_click)


# Entry point for command line usage

def main():
    import argparse

    parser = argparse.ArgumentParser(description='Research Paper Summarizer')
    parser.add_argument('action', choices=['summarize_pdfs'], help='Action to perform')
    parser.add_argument('--pdf_directory', required=True, help='Directory containing PDF files to summarize')
    parser.add_argument('--output_directory', required=True, help='Output directory for summaries')
    parser.add_argument('--prompt', required=True, help='Prompt for summarization')
    parser.add_argument('--model', default='llama3-8b-8192', choices=["llama3-8b-8192", "llama3-70b-8192", "gpt-4o-mini", "gpt-4o", "gpt-4"], help='Model to use for summarization')
    parser.add_argument('--openai_api_key', help='OpenAI API key (required for OpenAI models)')
    parser.add_argument('--groq_api_key', help='Groq API key (required for Groq models)')
    parser.add_argument('--voyageai_api_key', required=True, help='VoyageAI API key')
    parser.add_argument('--chunk_size', type=int, default=8000, help='Chunk size for text splitting')
    parser.add_argument('--chunk_overlap', type=int, default=500, help='Chunk overlap for text splitting')
    parser.add_argument('--similarity_threshold', type=float, default=0.8, help='Similarity threshold for filtering redundant chunks')
    parser.add_argument('--num_clusters', type=int, default=10, help='Number of clusters for K-means clustering')

    args = parser.parse_args()

    summarizer = ResSum()
    
    if args.action == 'summarize_pdfs':
        # Check if the required API key is provided based on the selected model
        if args.model in ["gpt-4o-mini", "gpt-4o", "gpt-4"] and not args.openai_api_key:
            parser.error(f"OpenAI API key is required for model {args.model}")
        elif args.model in ["llama3-8b-8192", "llama3-70b-8192"] and not args.groq_api_key:
            parser.error(f"Groq API key is required for model {args.model}")

        summarizer.summarize_pdfs(
            args.pdf_directory,
            args.output_directory,
            args.prompt,
            args.openai_api_key,
            args.groq_api_key,
            args.voyageai_api_key,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
            similarity_threshold=args.similarity_threshold,
            num_clusters=args.num_clusters
        )
    else:
        print(f"Unknown action: {args.action}")

if __name__ == '__main__':
    main()




