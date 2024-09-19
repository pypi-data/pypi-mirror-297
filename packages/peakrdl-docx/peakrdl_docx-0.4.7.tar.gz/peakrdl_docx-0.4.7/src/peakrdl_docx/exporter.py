"""PeakRDL Docx exporter."""

#__authors__ = ["Alex Nijmeijer <alex.nijmeijer at neads.nl>"]

from collections import OrderedDict
from dataclasses import dataclass
from functools import reduce
from pathlib import Path
from operator import mul
from typing import List, Optional, Union
import re
import os.path
import markdown

from systemrdl.messages import MessageHandler  # type: ignore # pylint: disable=import-error
from systemrdl.node import (  # type: ignore # pylint: disable=import-error
    AddressableNode,
    AddrmapNode,
    FieldNode,
    MemNode,
    Node,
    RegfileNode,
    RegNode,
    RootNode,
)

from pyquery import PyQuery as pq     # pylint: disable=import-error
from docx import Document             # pylint: disable=import-error
import docx                           # pylint: disable=import-error
from  .md2commands import md2commands  # pylint: disable=import-error

#from https://github.com/python-openxml/python-docx/issues/74
def _add_hyperlink(paragraph, url, text, color='0000FF', underline=False):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
    else :
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink




class DocxExporter:  # pylint: disable=too-few-public-methods
    """PeakRDL Docx exporter main class."""

    @dataclass
    class GenStageOutput:
        """Generation stage output."""

        node: Node
        """Node on which generation has been performed."""

        table_row: "OrderedDict[str, Union[str, int]]"
        """Row for the parent table."""

        generated: str
        """Docx generated during this stage."""

    def __init__ (self, **kwargs: 'Any') -> None:
        self.TableStyle = kwargs.pop("TableStyle", None)
        self.ParagraphStyle =  kwargs.pop("ParagraphStyle", None)
        self.ListStyle =  kwargs.pop("ListStyle", None)
        self.TitleStyle =  kwargs.pop("TitleStyle", None)

        self.title = "" # type: str

       

    @staticmethod
    def _heading(depth: int, title: str):
        """Generate Docx heading of a given depth with newline envelope.

        Arguments:
            depth -- heading depth (number of hashes)
            title -- heading title

        Returns:
            Formatted Docx heading.
        """
        return "\n" + "#" * depth + f" {title}\n"

    def _wordTable(self, data) :
        ret = ""

        ret += "\nTableHeader " + "".join(
            f"|{key}" for key, value in data[0].items()
            )+ "|"

        for i in data :
            ret += "\nTableData " + "".join(
            f"|{value}" for key, value in i.items()
            )+ "|\n"

        return ret

    @staticmethod
    def _addrnode_info(node: AddressableNode):
        """Generate AddressableNode basic information dictionary."""
        ret: "OrderedDict[str, str]" = OrderedDict()

        set_index = False
        if node.is_array and node.current_idx is None:
            node.current_idx = [0]
            set_index = True
        ret["Absolute Address"] = f"0x{node.absolute_address:X}"
        ret["Base Offset"] = f"0x{node.raw_address_offset:X}"
        if node.is_array and node.array_dimensions is not None and set_index:
            ret["Size"] = f"0x{node.size * reduce(mul, node.array_dimensions, 1):X}"
        else:
            ret["Size"] = f"0x{node.size:X}"

        if node.is_array:
            ret["Array Dimensions"] = str(node.array_dimensions)
            ret["Array Stride"] = f"0x{node.array_stride:X}"
            ret["Total Size"] = f"0x{node.total_size:X}"

        return ret

    def _addrnode_info_md(self, node: AddressableNode) -> str:
        """Generate AddressableNode basic information as a Docx list."""
        return "\n" + "".join(
            f"addrnode {key}: {value}\n" for key, value in self._addrnode_info(node).items()
        )

    @staticmethod
    def _node_name_sanitized(node: Node) -> str:
        """Get the Node name as HTML without newlines.

        Needed for proper inclusion in tables.
        """
        name = node.get_html_name()
        if name is None:
            name = "—"
        else:
            name = name.replace("\n", "")
        return name

    def _addrnode_header(
        self, node: AddressableNode, msg: MessageHandler, heading_level: int
    ) -> str:
        """Get the AddressableNode header.

        Arguments:
            node -- node to generate the header for.
            msg -- message handler from top-level.
            heading_level -- Docx heading level.
        """
        if isinstance(node, AddrmapNode):
            node_type_name = "address map"
        elif isinstance(node, RegfileNode):
            node_type_name = "register file"
        elif isinstance(node, MemNode):
            node_type_name = "memory"
        elif isinstance(node, RegNode):
            node_type_name = "register"
        else:
            node_type_name = "addressable node"
            msg.warning(f"Unsupported type of node ({node.__class__.__name__}).")

        ret = "heading " + str(heading_level) + f" {node.inst_name} {node_type_name}\n"

        ret += self._addrnode_info_md(node) # addrnode_info is bullet list

        desc = node.get_html_desc(markdown.Markdown() ) # this return a markdown-formatted string (with possible image references)
        ret += md2commands(desc)

        return ret

    def _addrnode_table_row(
        self, node: AddressableNode
    ) -> "OrderedDict[str, Union[str, int]]":
        """Generate AddressableNode table row.

        The "Offset" is an integer so that it can be formatted in the parent
        node.
        """
        offset = node.address_offset
        identifier = node.inst_name
        if node.is_array:
            assert node.current_idx is not None
            identifier += "".join(f"[{idx}]" for idx in node.current_idx)
        name = self._node_name_sanitized(node)

        table_row: "OrderedDict[str, Union[str, int]]" = OrderedDict()
        table_row["Offset"] = offset
        table_row["Identifier"] = identifier
        table_row["Name"] = name

        return table_row

    def _convert_to_word (self, output_path, cmdstr) :
      
        self.mydocument = Document()
        
        for s in cmdstr.splitlines() :
            splitted = s.strip().split(' ', 1) # remove spaces to left/right first
            if len(splitted) > 0 :
                command = splitted[0].strip()
            else :
                command = ""
            if len(splitted) > 1 :
                arg = splitted[1]
            else :
                arg =""

            i=0
            if len(splitted) > 0 :
                match command :

                    case "TableHeader" :
                        trimmed=arg.strip()

                        colcnt =  len(arg.split('|')) - 2

                        table = self.mydocument.add_table(rows=1, cols=colcnt)
                        table.style= self.TableStyle
                        hdr_cells = table.rows[0].cells
                        i=0
                        for k in arg.split('|') :
                            if 0 < i < colcnt+1:  # first col is dummy
                                hdr_cells[i-1].text=k
                            i=i+1

                    case "TableData" :
                        trimmed=arg.strip()
                        row_cells = table.add_row().cells
                        i=0
                        for v in trimmed.split('|') :
                            if 0 < i <= colcnt:  # first col is dummy
                                row_cells[i-1].text = v
                            i=i+1

                    case "para" :
                        paragraph = self.mydocument.add_paragraph(arg)
                        paragraph.style=self.ParagraphStyle

                    case "text" :
                        paragraph.add_run(arg+" ")

                    case "bolt" :
                        paragraph.add_run(arg+" ").bold = True

                    case "italic" :
                        paragraph.add_run(arg+" ").italic = True

                    case "image" :
                        file_path = re.match("\"(.+)\"", arg).group(1)
                        if os.path.isfile(file_path) :
                            self.mydocument.add_picture(file_path) # , width=Inches(1.25))
                        else :
                            paragraph = self.mydocument.add_paragraph(f"File not found {file_path}")
                            print("warning: image file not found:", file_path)

                    case "hlink" :
                        res = re.findall(r"\"(.+)\"(.*)", arg)
                        url=res[0][0]
                        text=res[0][1]
                        _add_hyperlink(paragraph, text=text, url=url)

                    case "heading" :
                        splitted = arg.strip().split(' ', 1) # remove spaces to left/right first
                        if len(arg) > 1 :
                            level = int(splitted[0].strip() )
                            arg   = splitted[1]
                            self.mydocument.add_heading(arg, level)

                    case "addrnode" :
                        self.mydocument.add_paragraph(arg, style=self.ListStyle)

                    case "title" :
                        self.mydocument.add_paragraph(arg, style=self.TitleStyle)
                    

        self.mydocument.save(output_path)

    def export(  # pylint: disable=too-many-arguments
        self,
        node: Union[AddrmapNode, RootNode],
        output_path: str,
        input_files: Optional[List[str]] = None,
        title: str = "",
        depth: int = 0,
    ):
        """Export the `node` to generated Python interface file.

        Arguments:
            node -- node to export.
            input_files -- list of input files.
            output_path -- path to the exported file.
            title -- Override title text.
            depth -- depth of generation (0 means all)
        """
        # Get the top node.
        top = node.top if isinstance(node, RootNode) else node

        if title != "" :
            gen = f"\ntitle {title}\n"
        else:
            gen = ""

        generated_from = ""
        if input_files is not None:
            generated_from += "\naddrnode " + "\naddrnode ".join(f for f in input_files)

        # Ensure proper format of the output path and that the directory exists.
        if not output_path.endswith(".docx"):
            raise ValueError("The output file is not Docx file.")
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        gen += f"para Do not edit. This document is generated from the following input documents: {generated_from}\n"
        gen += f"para Top level node: {top.inst_name}\n"

        # Run generation.
        gen += self._add_addrmap_regfile_mem(top, node.env.msg, depth - 1).generated
        self._convert_to_word(output_path, gen)

    def _add_addrmap_regfile_mem(
        self,
        node: Union[AddrmapNode, RegfileNode, MemNode],
        msg: MessageHandler,
        depth: int,
    ) -> GenStageOutput:
        """Generate addrmap, regfile or memory.

        Arguments:
            node -- MemNode, RegfileNode or AddrmapNode.
            msg -- message handler from top-level.
            depth -- depth of generation left.

        Keyword Arguments:
            is_top -- if the current not is the top node. If True the
                specification is embedded as class member.

        Returns:
            Generated addrmap output.
        """
        members: List[DocxExporter.GenStageOutput] = []
        member_gen: str = ""
        # Don't unroll register arrays when they are inside memories.
        # Memories can contain hundreds of entires.
        not_memory = not isinstance(node, MemNode)
        for child in node.children(unroll=not_memory, skip_not_present=False):
            if isinstance(child, (AddrmapNode, RegfileNode, MemNode)):
                output = self._add_addrmap_regfile_mem(child, msg, depth - 1)
                member_gen += output.generated
                members.append(output)

            elif isinstance(child, RegNode):
                output = self._add_reg(child, msg)
                member_gen += output.generated
                members.append(output)
            else:
                msg.warning(
                    f"Unsupported type of node ({child.__class__.__name__}) "
                    f"for {'/'.join(child.get_path_segments())}."
                )

        gen: str = self._addrnode_header(node, msg, 1)

        if len(members) == 0:
            gen += "para No supported members.\n"
        else:
            # Find the maximum width of the offset hex int and format the
            # offset for all members.
            base_addr_digits = max(
                map(lambda m: len(f'{m.table_row["Offset"]:X}'), members)
            )
            for member in members:
                member.table_row["Offset"] = (
                    f'0x{member.table_row["Offset"]:0{base_addr_digits}X}'
                )

            gen += "\n"
            gen += self._wordTable([*map(lambda m: m.table_row, members)])

        return DocxExporter.GenStageOutput(
            node,
            self._addrnode_table_row(node),
            gen + (member_gen if depth != 0 else ""),
        )

    def _add_reg(self, node: RegNode, msg: MessageHandler) -> GenStageOutput:
        """Generate register.

        Arguments:
            node -- RegNode.
            msg -- message handler from top-level.

        Returns:
            Generated register output.
        """
        field_gen: str = ""
        members: List[DocxExporter.GenStageOutput] = []
        for field in node.fields(skip_not_present=True):
            output = self._add_field(field, msg)
            field_gen += output.generated
            members.append(output)

        gen: str = self._addrnode_header(node, msg, 2)
        gen += self._wordTable([*map(lambda m: m.table_row, members)])
        gen += "\n"

        return DocxExporter.GenStageOutput(
            node, self._addrnode_table_row(node), gen + field_gen
        )

    def _add_field(
        self,
        node: FieldNode,
        msg: MessageHandler,  # pylint: disable=unused-argument
    ) -> GenStageOutput:
        """Generate field.

        Arguments:
            node -- FieldNode.
            msg -- message handler from top-level.

        Returns:
            Generated field output.
        """
        if node.msb == node.lsb:
            bits = str(node.msb)
        else:
            bits = f"{node.msb}:{node.lsb}"

        identifier = node.inst_name

        access = node.get_property("sw").name
        if node.get_property("onread") is not None:
            access += ", " + node.get_property("onread").name
        if node.get_property("onwrite") is not None:
            access += ", " + node.get_property("onwrite").name

        reset_value: str = node.get_property("reset", default="—")
        if isinstance(reset_value, int):
            reset = f"0x{reset_value:X}"
        else:
            reset = str(reset_value)

        name = self._node_name_sanitized(node)

        table_row: "OrderedDict[str, Union[str, int]]" = OrderedDict()
        table_row["Bits"] = bits
        table_row["Identifier"] = identifier
        table_row["Access"] = access
        table_row["Reset"] = reset
        table_row["Name"] = name

        gen = ""
        desc = node.get_html_desc()
        if desc is not None:
            gen = f"heading 3 {node.inst_name} field\n " + "para " + pq(desc).text() + "\n"

        return DocxExporter.GenStageOutput(node, table_row, gen)
